"""Export meeting notes to whatever format the reader needs.

Markdown is what gets stored, because it survives everything and stays
readable in a plain editor. The rest are conversions produced on demand:

    .md    the stored file, copied as-is
    .txt   markdown stripped, for pasting into anything
    .html  self-contained, opens in a browser and prints cleanly
    .docx  for people who live in Word
    .pdf   for sending on, where layout must not shift

The markdown here is deliberately a small subset — headings, bullets,
bold — because that is all write_notes ever produces. A full markdown
engine would be a dependency with nothing to do.
"""
from __future__ import annotations

import html
import logging
import os
import re
import shutil

log = logging.getLogger("talkwithme.export")

FORMATS = (
    ("Markdown", ".md"),
    ("Tekstbestand", ".txt"),
    ("Webpagina", ".html"),
    ("Word-document", ".docx"),
    ("PDF", ".pdf"),
)

_BOLD = re.compile(r"\*\*(.+?)\*\*")
_ITALIC = re.compile(r"(?<!\*)\*(?!\s)(.+?)(?<!\s)\*(?!\*)")
_CODE = re.compile(r"`([^`]+)`")


# ---- parsing ---------------------------------------------------------

def _blocks(markdown: str):
    """Yield (kind, text, level) for each line we care about."""
    for raw in markdown.split("\n"):
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            yield ("blank", "", 0)
        elif stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            yield ("heading", stripped[level:].strip(), min(level, 4))
        elif stripped in ("---", "***", "___"):
            yield ("rule", "", 0)
        elif stripped.startswith(("- ", "* ", "+ ")):
            yield ("bullet", stripped[2:].strip(), 0)
        elif re.match(r"^\d+[.)]\s", stripped):
            yield ("number", re.sub(r"^\d+[.)]\s+", "", stripped), 0)
        else:
            yield ("text", stripped, 0)


def _plain(text: str) -> str:
    text = _BOLD.sub(r"\1", text)
    text = _ITALIC.sub(r"\1", text)
    return _CODE.sub(r"\1", text)


# ---- converters ------------------------------------------------------

def to_text(markdown: str) -> str:
    out: list[str] = []
    for kind, text, level in _blocks(markdown):
        if kind == "blank":
            out.append("")
        elif kind == "heading":
            plain = _plain(text)
            out.extend(["", plain.upper() if level <= 2 else plain])
            if level <= 2:
                out.append("-" * len(plain))
        elif kind == "rule":
            out.append("-" * 60)
        elif kind == "bullet":
            out.append(f"  - {_plain(text)}")
        elif kind == "number":
            out.append(f"  {_plain(text)}")
        else:
            out.append(_plain(text))
    # Collapse the runs of blanks the headings introduce.
    cleaned: list[str] = []
    for line in out:
        if line == "" and cleaned[-1:] == [""]:
            continue
        cleaned.append(line)
    return "\n".join(cleaned).strip() + "\n"


def _inline_html(text: str) -> str:
    escaped = html.escape(text)
    escaped = _BOLD.sub(r"<strong>\1</strong>", escaped)
    escaped = _ITALIC.sub(r"<em>\1</em>", escaped)
    return _CODE.sub(r"<code>\1</code>", escaped)


def to_html(markdown: str, title: str = "Vergadernotities") -> str:
    body: list[str] = []
    list_open: str | None = None

    def close_list():
        nonlocal list_open
        if list_open:
            body.append(f"</{list_open}>")
            list_open = None

    for kind, text, level in _blocks(markdown):
        if kind in ("bullet", "number"):
            wanted = "ul" if kind == "bullet" else "ol"
            if list_open != wanted:
                close_list()
                body.append(f"<{wanted}>")
                list_open = wanted
            body.append(f"<li>{_inline_html(text)}</li>")
            continue
        close_list()
        if kind == "heading":
            body.append(f"<h{level}>{_inline_html(text)}</h{level}>")
        elif kind == "rule":
            body.append("<hr>")
        elif kind == "text":
            body.append(f"<p>{_inline_html(text)}</p>")
    close_list()

    return f"""<!doctype html>
<html lang="nl">
<meta charset="utf-8">
<title>{html.escape(title)}</title>
<style>
  body {{ max-width: 46em; margin: 3em auto; padding: 0 1.5em;
         font: 16px/1.65 "Segoe UI", system-ui, sans-serif; color: #0A2540; }}
  h1 {{ font-size: 1.7em; margin-bottom: .2em; }}
  h2 {{ font-size: 1.15em; margin-top: 2em; padding-bottom: .3em;
        border-bottom: 1px solid #E3E8EE; }}
  li {{ margin: .3em 0; }}
  hr {{ border: 0; border-top: 1px solid #E3E8EE; margin: 2.5em 0; }}
  code {{ background: #F6F9FC; padding: .1em .35em; border-radius: 3px; }}
  @media print {{ body {{ margin: 0; max-width: none; }} }}
</style>
{os.linesep.join(body)}
</html>
"""


def to_docx(markdown: str, path: str, title: str = "Vergadernotities") -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.core_properties.title = title

    for kind, text, level in _blocks(markdown):
        if kind == "heading":
            doc.add_heading(_plain(text), level=min(level, 4))
        elif kind == "bullet":
            _add_runs(doc.add_paragraph(style="List Bullet"), text)
        elif kind == "number":
            _add_runs(doc.add_paragraph(style="List Number"), text)
        elif kind == "text":
            _add_runs(doc.add_paragraph(), text)
        elif kind == "rule":
            para = doc.add_paragraph()
            para.add_run("―" * 30).font.size = Pt(8)

    doc.save(path)


def _add_runs(paragraph, text: str) -> None:
    """Split on **bold** so the emphasis survives into Word."""
    for index, part in enumerate(_BOLD.split(text)):
        if not part:
            continue
        run = paragraph.add_run(_plain(part))
        run.bold = index % 2 == 1


def to_pdf(markdown: str, path: str, title: str = "Vergadernotities") -> None:
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import (HRFlowable, ListFlowable, ListItem,
                                     Paragraph, SimpleDocTemplate, Spacer)

    styles = getSampleStyleSheet()
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontSize=10.5,
                           leading=15.5, spaceAfter=7, alignment=TA_LEFT)
    heads = {
        1: ParagraphStyle("H1", parent=styles["Heading1"], fontSize=17, spaceAfter=10),
        2: ParagraphStyle("H2", parent=styles["Heading2"], fontSize=12.5,
                           spaceBefore=16, spaceAfter=7),
        3: ParagraphStyle("H3", parent=styles["Heading3"], fontSize=11, spaceBefore=12),
        4: ParagraphStyle("H4", parent=styles["Heading4"], fontSize=10.5),
    }

    story: list = []
    pending: list = []
    pending_kind: str | None = None

    def flush_list():
        nonlocal pending, pending_kind
        if pending:
            story.append(ListFlowable(
                pending, bulletType="bullet" if pending_kind == "bullet" else "1",
                leftIndent=14, bulletFontSize=8))
            story.append(Spacer(1, 5))
        pending, pending_kind = [], None

    for kind, text, level in _blocks(markdown):
        if kind in ("bullet", "number"):
            if pending_kind and pending_kind != kind:
                flush_list()
            pending_kind = kind
            pending.append(ListItem(Paragraph(_inline_html(text), body), leftIndent=14))
            continue
        flush_list()
        if kind == "heading":
            story.append(Paragraph(_inline_html(text), heads[level]))
        elif kind == "rule":
            story.append(Spacer(1, 8))
            story.append(HRFlowable(width="100%", color="#CFD7DF"))
            story.append(Spacer(1, 8))
        elif kind == "text":
            story.append(Paragraph(_inline_html(text), body))
    flush_list()

    SimpleDocTemplate(path, pagesize=A4, title=title,
                       leftMargin=2.2 * cm, rightMargin=2.2 * cm,
                       topMargin=2 * cm, bottomMargin=2 * cm).build(story)


# ---- dispatch --------------------------------------------------------

def export(markdown: str, target_path: str, source_path: str | None = None,
           title: str = "Vergadernotities") -> str:
    """Write `markdown` to target_path in the format its extension implies."""
    extension = os.path.splitext(target_path)[1].lower()

    if extension == ".md":
        if source_path and os.path.exists(source_path):
            shutil.copyfile(source_path, target_path)
        else:
            _write(target_path, markdown)
    elif extension == ".txt":
        _write(target_path, to_text(markdown))
    elif extension == ".html":
        _write(target_path, to_html(markdown, title))
    elif extension == ".docx":
        to_docx(markdown, target_path, title)
    elif extension == ".pdf":
        to_pdf(markdown, target_path, title)
    else:
        raise ValueError(f"Onbekend formaat: {extension or 'geen extensie'}")

    log.info("geëxporteerd naar %s", target_path)
    return target_path


def _write(path: str, content: str) -> None:
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)
